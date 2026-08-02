"""
Auto-extracts text from an uploaded policy file (PDF or DOCX) so the admin
never has to retype the content by hand.

Returns (extracted_text, note) — note is empty on success, or a short
explanation if extraction failed / the format isn't supported.
"""


def extract_text_from_file(django_file):
    if not django_file:
        return '', ''

    name = django_file.name.lower()

    try:
        if name.endswith('.pdf'):
            return _extract_pdf(django_file)
        elif name.endswith('.docx'):
            return _extract_docx(django_file)
        elif name.endswith('.doc'):
            return '', "Old .doc format isn't supported for auto-extraction — please re-save as .docx or .pdf, or add notes manually below."
        elif name.endswith('.txt'):
            return _extract_txt(django_file)
        else:
            return '', 'Unsupported file type for auto text extraction. Supported: PDF, DOCX, TXT.'
    except Exception as e:
        return '', f'Could not extract text automatically ({e}). The file is still saved and downloadable.'


def _extract_pdf(django_file):
    try:
        from pypdf import PdfReader
    except ImportError:
        return '', "PDF extraction needs the 'pypdf' package. Run: pip install pypdf"

    django_file.seek(0)
    reader = PdfReader(django_file)
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ''
        if text.strip():
            pages.append(text.strip())
    full_text = '\n\n'.join(pages).strip()

    if not full_text:
        return '', 'No selectable text found in this PDF (it may be a scanned image). The file is still saved and downloadable.'
    return full_text, ''


def _extract_docx(django_file):
    try:
        import docx
    except ImportError:
        return '', "DOCX extraction needs the 'python-docx' package. Run: pip install python-docx"

    django_file.seek(0)
    document = docx.Document(django_file)
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also pull text out of any tables in the doc
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(' | '.join(cells))

    full_text = '\n\n'.join(parts).strip()
    if not full_text:
        return '', 'No text found in this document.'
    return full_text, ''


def _extract_txt(django_file):
    django_file.seek(0)
    raw = django_file.read()
    try:
        text = raw.decode('utf-8')
    except UnicodeDecodeError:
        text = raw.decode('latin-1', errors='ignore')
    return text.strip(), ''